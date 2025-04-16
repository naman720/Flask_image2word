from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
import os
from werkzeug.utils import secure_filename
import uuid
from PIL import Image

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

SUPPORTED_FORMATS = ['.jpg', '.jpeg', '.png', '.bmp', '.gif']

def convert_to_supported_format(path):
    with Image.open(path) as img:
        rgb_im = img.convert('RGB')
        new_path = path.rsplit('.', 1)[0] + '.jpg'
        rgb_im.save(new_path, format='JPEG')
        return new_path

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        images = request.files.getlist('images')
        descriptions = request.form.getlist('descriptions')

        image_width = float(request.form.get('image_width', 5))
        image_height = float(request.form.get('image_height', 5))

        margin_top = float(request.form.get('margin_top', 1))
        margin_bottom = float(request.form.get('margin_bottom', 1))
        margin_left = float(request.form.get('margin_left', 1))
        margin_right = float(request.form.get('margin_right', 1))

        doc = Document()
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(margin_top)
            section.bottom_margin = Inches(margin_bottom)
            section.left_margin = Inches(margin_left)
            section.right_margin = Inches(margin_right)

        table = doc.add_table(rows=0, cols=2)
        table.autofit = False

        for i in range(0, len(images), 2):
            row = table.add_row().cells

            for j in range(2):
                if i + j < len(images):
                    image = images[i + j]
                    desc = descriptions[i + j]

                    filename = f"{uuid.uuid4()}_{secure_filename(image.filename)}"
                    path = os.path.join(UPLOAD_FOLDER, filename)
                    image.save(path)

                    ext = os.path.splitext(path)[1].lower()
                    if ext not in SUPPORTED_FORMATS:
                        path = convert_to_supported_format(path)

                    paragraph = row[j].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(path, width=Inches(image_width), height=Inches(image_height))

                    row[j].add_paragraph(desc)

        output_path = os.path.join(UPLOAD_FOLDER, 'output.docx')
        doc.save(output_path)
        return send_file(output_path, as_attachment=True)

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)